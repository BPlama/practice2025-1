# Логика для редактора текста

import tkinter as tk
from tkinter import filedialog
from docx import Document
from docx.shared import Pt


def open_file(text_area, word_count_callback):
    file_path = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt")])
    if file_path:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            text_area.delete(1.0, tk.END)
            text_area.insert(tk.END, content)
            word_count_callback()


def save_txt_file(text_area):
    file_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt")])
    if file_path:
        with open(file_path, 'w', encoding='utf-8') as file:
            content = text_area.get(1.0, tk.END)
            file.write(content)


def save_docx_file(text_area, font_size):
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
    if not file_path:
        return

    doc = Document()
    index = "1.0"
    while True:
        if index == text_area.index(tk.END):
            break

        line_end = index.split('.')[0] + ".end"
        line_text = text_area.get(index, line_end)

        if line_text.strip() == "":
            doc.add_paragraph("")
            index = text_area.index(f"{line_end}+1c")
            continue

        p = doc.add_paragraph()
        i = 0
        while i < len(line_text):
            char_index = f"{index}+{i}c"
            char = line_text[i]
            tags = text_area.tag_names(char_index)

            run = p.add_run(char)
            font = run.font
            font.size = Pt(font_size)
            font.bold = 'bold' in tags
            font.italic = 'italic' in tags
            font.underline = 'underline' in tags

            i += 1

        index = text_area.index(f"{line_end}+1c")

    doc.save(file_path)


def find_text(text_area, search_entry):
    text_area.tag_remove('highlight', '1.0', tk.END)
    search_word = search_entry.get()
    if search_word:
        start_pos = '1.0'
        while True:
            start_pos = text_area.search(search_word, start_pos, stopindex=tk.END)
            if not start_pos:
                break
            end_pos = f"{start_pos}+{len(search_word)}c"
            text_area.tag_add('highlight', start_pos, end_pos)
            text_area.tag_config('highlight', background='lightblue')
            start_pos = end_pos


def toggle_format(text_area, tag):
    try:
        current_tags = text_area.tag_names("sel.first")
        if tag in current_tags:
            text_area.tag_remove(tag, "sel.first", "sel.last")
        else:
            text_area.tag_add(tag, "sel.first", "sel.last")
    except tk.TclError:
        pass


def update_word_count(text_area, word_count_var):
    content = text_area.get("1.0", "end-1c")
    words = content.split()
    word_count_var.set(f"\u0421\u043b\u043e\u0432: {len(words)}")


def change_font_size(text_area, font_size):
    size = font_size.get()
    text_area.config(font=("Arial", size))
    for tag in ["bold", "italic", "underline"]:
        style = []
        if tag == "bold":
            style = ("Arial", size, "bold")
        elif tag == "italic":
            style = ("Arial", size, "italic")
        elif tag == "underline":
            style = ("Arial", size, "underline")
        text_area.tag_configure(tag, font=style)
